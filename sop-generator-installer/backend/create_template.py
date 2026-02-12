"""
Script to create/update the SOP Word template with correct formatting
Follows Guideline V2 specifications exactly
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_table_borders(table):
    """Add borders to table following guideline specs"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Add table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_shading(cell, shade_pct=None):
    """Set cell shading (15% darker for gateways)"""
    tcPr = cell._element.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:shd')

    if shade_pct == 15:
        # 15% darker (light gray)
        tcVAlign.set(qn('w:fill'), 'D9D9D9')
    else:
        # Default/white
        tcVAlign.set(qn('w:fill'), 'FFFFFF')

    tcPr.append(tcVAlign)

def set_font(run, name='Avenir LT Std 45 Book', size=11, bold=False):
    """Set font properties according to guideline"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    # Also set for complex scripts
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def create_sop_template(output_path='final_master_template.docx'):
    """
    Create a complete SOP template following Guideline V2
    Uses Jinja2 templating syntax for docxtpl
    """

    doc = Document()

    # Set default font for document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Avenir LT Std 45 Book'
    font.size = Pt(11)

    # ===== HEADER SECTION =====
    # Add logo placeholder (user should replace manually or via code)
    # header = doc.sections[0].header
    # You can add image to header here if needed

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('{{ process_name }}')
    set_font(title_run, size=16, bold=True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # ===== METADATA TABLE =====
    metadata_table = doc.add_table(rows=2, cols=4)
    metadata_table.style = 'Light Grid Accent 1'

    # Row 1
    metadata_table.cell(0, 0).text = 'Process Name:'
    metadata_table.cell(0, 1).text = '{{ process_name }}'
    metadata_table.cell(0, 2).text = 'Process Code:'
    metadata_table.cell(0, 3).text = '{{ process_code }}'

    # Row 2
    metadata_table.cell(1, 0).text = 'Issued By:'
    metadata_table.cell(1, 1).text = '{{ issued_by }}'
    metadata_table.cell(1, 2).text = 'Released Date:'
    metadata_table.cell(1, 3).text = '{{ release_date }}'

    doc.add_paragraph()

    # ===== CONTENT SECTIONS =====
    sections = [
        ('Purpose', '{{ purpose }}'),
        ('Scope', '{{ scope }}'),
        ('Abbreviations and Definitions', '{{ abbreviations }}'),
        ('Referenced Documents and Approvals', '{{ references }}'),
        ('Key Process Inputs', '{{ inputs }}'),
        ('Key Process Outputs', '{{ outputs }}')
    ]

    for section_title, section_content in sections:
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section_title)
        set_font(heading_run, size=12, bold=True)

        content = doc.add_paragraph(section_content)
        set_font(content.runs[0], size=11)
        doc.add_paragraph()

    # ===== PROCESS DESCRIPTION TABLE =====
    heading = doc.add_paragraph()
    heading_run = heading.add_run('Process Description')
    set_font(heading_run, size=14, bold=True)
    doc.add_paragraph()

    # Create main table with exact column widths from guideline
    # Total width: 10.12", Ref: 0.52", Process: 5", R/A/C/I: 1" each, SLA: 0.62"
    table = doc.add_table(rows=1, cols=7)
    table.autofit = False
    table.allow_autofit = False

    # Set column widths (in inches)
    widths = [0.52, 5.0, 1.0, 1.0, 1.0, 1.0, 0.62]
    for i, width in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = Inches(width)

    # Header row
    header_cells = table.rows[0].cells
    headers = ['Ref.', 'Process Description', 'R', 'A', 'C', 'I', 'SLA']

    for i, header_text in enumerate(headers):
        cell = header_cells[i]
        cell.text = header_text
        # Format header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, size=11, bold=True)
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        # Blue header background
        set_cell_shading(cell, shade_pct=None)
        tcPr = cell._element.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:shd')
        tcVAlign.set(qn('w:fill'), '4472C4')  # Blue
        tcPr.append(tcVAlign)

    # Add docxtpl template row using {%tr for step in steps %}
    # This will create one row per step

    # First add the opening tag row
    opening_row = table.add_row()
    opening_row.cells[0].text = '{%tr for step in steps %}'

    # Now add the actual template row with data
    row_tmpl = table.add_row()
    cells = row_tmpl.cells

    # Ref cell
    cells[0].text = '{{step.ref}}'
    for paragraph in cells[0].paragraphs:
        for run in paragraph.runs:
            set_font(run, size=12, bold=True)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Process Description cell
    cells[1].text = '{{step.desc}}'
    for paragraph in cells[1].paragraphs:
        for run in paragraph.runs:
            set_font(run, size=11, bold=False)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # R, A, C, I, SLA columns (empty for user to fill)
    for i in range(2, 7):
        cells[i].text = ''

    # Close the loop after the row
    closing_row = table.add_row()
    closing_row.cells[0].text = '{%tr endfor %}'

    add_table_borders(table)

    # Save template
    doc.save(output_path)
    print(f"[OK] Template created successfully: {output_path}")
    print("[NOTE] You need to manually add your logo to the header")
    print("[NOTE] The template uses Jinja2 syntax and works with docxtpl")
    print("\n[IMPORTANT] Gateway row shading must be applied dynamically in code")
    print("            (docxtpl doesn't support conditional formatting natively)")

if __name__ == '__main__':
    create_sop_template()
