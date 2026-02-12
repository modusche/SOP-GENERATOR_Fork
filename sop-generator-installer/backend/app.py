import io
import re
import threading
import webbrowser
import os
import sys
from flask import Flask, render_template, request, send_file, make_response, jsonify, session
from docxtpl import DocxTemplate
from waitress import serve

# --- Debug logging to file ---
def debug_log(msg):
    log_path = os.path.join(os.environ.get('LOCALAPPDATA', '.'), 'SOP_Generator', 'debug.log')
    os.makedirs(os.path.dirname(log_path), exist_ok=True)
    with open(log_path, 'a') as f:
        f.write(f"{msg}\n")

debug_log(f"=== App starting ===")
debug_log(f"sys.frozen: {getattr(sys, 'frozen', False)}")
debug_log(f"sys.executable: {sys.executable}")
debug_log(f"os.getcwd(): {os.getcwd()}")


# Import our custom BPMN parser
from bpmn_parser import parse_bpmn_to_sop, extract_metadata_from_bpmn

# Import history manager
from history_manager import HistoryManager

# Import archive manager
from archive_manager import ArchiveManager

# --- Helper function to find bundled files ---
def resource_path(relative_path):
    # Check for Nuitka (sets __compiled__ at module level) or PyInstaller (sets sys.frozen)
    is_nuitka = "__compiled__" in globals()
    is_pyinstaller = getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS')

    if is_pyinstaller:
        base_path = sys._MEIPASS
        debug_log(f"Using PyInstaller _MEIPASS: {base_path}")
    elif is_nuitka or 'Temp' in sys.executable:
        # Nuitka onefile extracts to temp - use executable's directory
        base_path = os.path.dirname(sys.executable)
        debug_log(f"Using Nuitka exe dir: {base_path}")
    else:
        # Running from source: use directory of this file so templates are always from backend/
        base_path = os.path.dirname(os.path.abspath(__file__))
        debug_log(f"Using source dir (app.py location): {base_path}")
        # Special handling for template file
        if relative_path == 'final_master_template_2.docx':
            deployment_path = os.path.join(base_path, 'SOP_Generator_Deployment', relative_path)
            if os.path.exists(deployment_path):
                return deployment_path
    result = os.path.join(base_path, relative_path)
    debug_log(f"resource_path('{relative_path}') -> {result} (exists: {os.path.exists(result)})")
    return result

# --- Allowed SOP templates (user can choose in UI) ---
ALLOWED_SOP_TEMPLATES = [
    ('final_master_template_2.docx', 'Default (Master)'),
    ('sabah_template.docx', 'Sabah'),
    ('sana_template.docx', 'Sana'),
    ('tarabut_template.docx', 'Tarabut'),
    ('window_world_template.docx', 'Window World'),
]
DEFAULT_SOP_TEMPLATE = 'final_master_template_2.docx'

# --- Flask App Initialization ---
template_folder = resource_path('templates')
app = Flask(__name__, template_folder=template_folder)
app.secret_key = 'sop-generator-secret-key-change-in-production'  # For session management

# --- App Data Directory (store runtime data in AppData, not exe folder) ---
def get_app_data_dir():
    if os.name == 'nt':  # Windows
        base = os.environ.get('LOCALAPPDATA', os.path.expanduser('~'))
    else:  # Linux/Mac
        base = os.path.join(os.path.expanduser('~'), '.config')
    app_dir = os.path.join(base, 'SOP_Generator')
    os.makedirs(app_dir, exist_ok=True)
    return app_dir

APP_DATA_DIR = get_app_data_dir()

# --- History Manager ---
history_manager = HistoryManager(os.path.join(APP_DATA_DIR, 'history'))

# --- Archive Manager ---
archive_manager = ArchiveManager(os.path.join(APP_DATA_DIR, 'archives'), os.path.join(APP_DATA_DIR, 'archive.db'))

# --- Core Logic using our custom BPMN parser ---
def parse_bpmn_to_context(xml_content, metadata):
    """
    Wrapper function that calls our comprehensive BPMN parser
    """
    return parse_bpmn_to_sop(xml_content, metadata)


def create_word_doc_from_template(context, template_name=None):
    """
    Create Word document with multi-paragraph structure and precise formatting.
    template_name: filename from ALLOWED_SOP_TEMPLATES (e.g. 'final_master_template_2.docx').
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        allowed = [t[0] for t in ALLOWED_SOP_TEMPLATES]
        if not template_name or template_name not in allowed:
            template_name = DEFAULT_SOP_TEMPLATE
        template_path = resource_path(template_name)

        # First, render metadata using docxtpl (for any {{variables}} in headers/etc)
        doc_template = DocxTemplate(template_path)

        # Create a simple context for metadata rendering (excluding steps)
        metadata_context = {k: v for k, v in context.items() if k != 'steps'}
        doc_template.render(metadata_context)

        # Now work with the rendered document using python-docx
        # Save to memory and reload as Document for easier manipulation
        temp_stream = io.BytesIO()
        doc_template.save(temp_stream)
        temp_stream.seek(0)
        doc = Document(temp_stream)

        # Get the tables
        if not doc.tables:
            raise Exception("No tables found in template")

        # Fix font for all front matter tables to size 12
        # Table 0: Header (Process Name, Code, etc.)
        # Table 1: Purpose
        # Table 2: Scope
        # Table 3: Abbreviations and Definitions
        # Table 4: Referenced Documents and Approvals
        # Table 5: Key Process Inputs/Outputs
        for table_idx in [0, 1, 2, 3, 4, 5]:
            if len(doc.tables) > table_idx:
                table = doc.tables[table_idx]
                # Iterate through all cells in the table
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Avenir LT Std 45 Book'
                                run.font.size = Pt(12)

        # Populate Table 4: Abbreviations and Definitions (index 3)
        if len(doc.tables) > 3:
            abbrev_table = doc.tables[3]
            abbreviations = context.get('abbreviations_list', [])

            # Clear existing rows (keep row 0=title and row 1=headers)
            rows_to_delete = list(range(len(abbrev_table.rows) - 1, 1, -1))
            for row_idx in rows_to_delete:
                abbrev_table._element.remove(abbrev_table.rows[row_idx]._element)

            # Add rows for each abbreviation
            if abbreviations:
                for abbrev in abbreviations:
                    row = abbrev_table.add_row()
                    row.cells[0].text = abbrev.get('term', '')
                    row.cells[1].text = abbrev.get('definition', '')
                    # Format cells - Terms column (column 0) should be bold
                    for idx, cell in enumerate(row.cells):
                        para = cell.paragraphs[0]
                        if para.runs:
                            para.runs[0].font.name = 'Avenir LT Std 45 Book'
                            para.runs[0].font.size = Pt(12)  # Changed from 11 to 12
                            # Make Terms column (column 0) bold
                            if idx == 0:
                                para.runs[0].font.bold = True
            else:
                # Add one empty row if no abbreviations provided
                row = abbrev_table.add_row()
                row.cells[0].text = 'N/A'
                row.cells[1].text = 'N/A'
                # Format cells
                for cell in row.cells:
                    para = cell.paragraphs[0]
                    if para.runs:
                        para.runs[0].font.name = 'Avenir LT Std 45 Book'
                        para.runs[0].font.size = Pt(12)

        # Populate Table 5: Referenced Documents and Approvals (index 4)
        if len(doc.tables) > 4:
            ref_table = doc.tables[4]
            references = context.get('references_list', [])

            # Clear existing rows (keep row 0=title and row 1=headers)
            rows_to_delete = list(range(len(ref_table.rows) - 1, 1, -1))
            for row_idx in rows_to_delete:
                ref_table._element.remove(ref_table.rows[row_idx]._element)

            # Add rows for each reference
            if references:
                for ref in references:
                    row = ref_table.add_row()
                    row.cells[0].text = ref.get('id', '')
                    row.cells[1].text = ref.get('title', '')
                    # Format cells - Document ID column (column 0) should be bold
                    for idx, cell in enumerate(row.cells):
                        para = cell.paragraphs[0]
                        if para.runs:
                            para.runs[0].font.name = 'Avenir LT Std 45 Book'
                            para.runs[0].font.size = Pt(12)  # Changed from 11 to 12
                            # Make Document ID column (column 0) bold
                            if idx == 0:
                                para.runs[0].font.bold = True
            else:
                # Add one empty row if no references provided
                row = ref_table.add_row()
                row.cells[0].text = 'N/A'
                row.cells[1].text = 'N/A'
                # Format cells
                for cell in row.cells:
                    para = cell.paragraphs[0]
                    if para.runs:
                        para.runs[0].font.name = 'Avenir LT Std 45 Book'
                        para.runs[0].font.size = Pt(12)

        # Get the process description table (Table 6 - Table 7 is General Policies)
        table = doc.tables[6]

        # Clear all rows except header (row 0)
        rows_to_delete = list(range(len(table.rows) - 1, 0, -1))  # Delete in reverse order
        for row_idx in rows_to_delete:
            table._element.remove(table.rows[row_idx]._element)

        # Now add rows for each step with proper formatting
        steps = context.get('steps', [])

        for step in steps:
            # Add new row
            new_row = table.add_row()
            cells = new_row.cells

            # --- Cell 0: Ref number ---
            ref_cell = cells[0]
            ref_cell.text = ''  # Clear any existing text
            ref_para = ref_cell.paragraphs[0]
            ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if step['ref']:  # Only add ref if it exists
                run = ref_para.add_run(step['ref'])
                run.font.name = 'Avenir LT Std 45 Book'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)  # RED

            # --- Cell 1: Process Description (multi-paragraph) ---
            desc_cell = cells[1]
            desc_cell.text = ''  # Clear any existing text

            # Remove default paragraph
            if desc_cell.paragraphs:
                p = desc_cell.paragraphs[0]
                p._element.getparent().remove(p._element)

            # Add each paragraph with specific formatting
            for para_data in step['paragraphs']:
                para = desc_cell.add_paragraph()

                # Set alignment
                if para_data['alignment'] == 'CENTER':
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif para_data['alignment'] == 'JUSTIFY':
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Add text with formatting
                if para_data['text']:  # Only add run if there's text
                    run = para.add_run(para_data['text'])
                    run.font.name = 'Avenir LT Std 45 Book'
                    run.font.size = Pt(para_data['font_size'])
                    run.font.bold = para_data['bold']
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black

            # --- Cells 2-6: RACI + SLA ---
            # All RACI fields must be Avenir LT Std 45 Book font size 9
            raci = step.get('raci', {})
            raci_map = {2: 'responsible', 3: 'accountable', 4: 'consulted', 5: 'informed'}
            for i in range(2, 7):
                # Clear cell and get paragraph
                cells[i].text = ''
                para = cells[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Set paragraph style font (this ensures the cell has the right font even when empty)
                para.style.font.name = 'Avenir LT Std 45 Book'
                para.style.font.size = Pt(9)

                if i == 6:  # SLA column - leave blank (handled by SLA merge logic)
                    run = para.add_run('')
                    run.font.name = 'Avenir LT Std 45 Book'
                    run.font.size = Pt(9)
                elif i in raci_map:  # R, A, C, I columns - use lane RACI values
                    value = raci.get(raci_map[i], 'N/A') or 'N/A'
                    run = para.add_run(value)
                    run.font.name = 'Avenir LT Std 45 Book'
                    run.font.size = Pt(9)

            # --- Apply gateway shading if needed ---
            if step.get('is_gateway', False):
                for cell_idx, cell in enumerate(new_row.cells):
                    if cell_idx == 6:
                        continue  # Skip SLA column - handled separately
                    tcPr = cell._element.get_or_add_tcPr()
                    # Remove existing shading
                    for shd in tcPr.findall(qn('w:shd')):
                        tcPr.remove(shd)
                    # Apply D9D9D9 shading
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'D9D9D9')  # 15% darker gray
                    shd.set(qn('w:val'), 'clear')
                    tcPr.append(shd)

        # --- SLA shading and vertical merging ---
        # Compute merge ranges: (start_step_idx, end_step_idx, sla_value)
        sla_merges = []
        idx = 0
        while idx < len(steps):
            step_item = steps[idx]
            sla = step_item.get('sla')
            sla_group = step_item.get('sla_group')

            if sla and not sla_group:
                # Task with own SLA - include following gateway cases
                merge_start = idx
                merge_end = idx
                j = idx + 1
                while j < len(steps) and steps[j].get('is_gateway', False):
                    merge_end = j
                    j += 1
                sla_merges.append((merge_start, merge_end, sla))
                idx = j
            elif sla_group:
                # Task in SLA group - include all group members + their gateway cases
                merge_start = idx
                merge_end = idx
                group_sla = sla
                j = idx + 1
                # Include gateway cases of this task
                while j < len(steps) and steps[j].get('is_gateway', False):
                    merge_end = j
                    j += 1
                # Continue with next tasks in same group
                while j < len(steps) and steps[j].get('sla_group') == sla_group:
                    merge_end = j
                    j += 1
                    # Include their gateway cases too
                    while j < len(steps) and steps[j].get('is_gateway', False):
                        merge_end = j
                        j += 1
                sla_merges.append((merge_start, merge_end, group_sla))
                idx = j
            else:
                # No SLA - skip this step and any following gateway cases
                idx += 1
                while idx < len(steps) and steps[idx].get('is_gateway', False):
                    idx += 1

        # Apply SLA shading and vertical merge
        for merge_start, merge_end, sla_value in sla_merges:
            for row_offset in range(merge_start, merge_end + 1):
                table_row_idx = row_offset + 1  # Row 0 is header
                if table_row_idx >= len(table.rows):
                    break
                sla_cell = table.rows[table_row_idx].cells[6]
                tcPr = sla_cell._element.get_or_add_tcPr()

                # Apply F2F2F2 shading (White, Background 1, Darker 5%)
                for existing_shd in tcPr.findall(qn('w:shd')):
                    tcPr.remove(existing_shd)
                shd_elem = OxmlElement('w:shd')
                shd_elem.set(qn('w:fill'), 'F2F2F2')
                shd_elem.set(qn('w:val'), 'clear')
                tcPr.append(shd_elem)

                # Vertical merge if multiple rows in range
                if merge_end > merge_start:
                    vMerge = OxmlElement('w:vMerge')
                    if row_offset == merge_start:
                        vMerge.set(qn('w:val'), 'restart')
                    tcPr.append(vMerge)

                # Write SLA value only in first row of merge
                if row_offset == merge_start:
                    p = sla_cell.paragraphs[0]
                    if p.runs:
                        p.runs[0].text = sla_value

        # --- Populate Table 7: General Policies ---
        if len(doc.tables) > 7:
            policies_table = doc.tables[7]
            policies = context.get('general_policies_list', [])

            # Clear existing data rows (keep row 0=headers only)
            rows_to_delete = list(range(len(policies_table.rows) - 1, 0, -1))
            for row_idx in rows_to_delete:
                policies_table._element.remove(policies_table.rows[row_idx]._element)

            if policies:
                for idx, policy in enumerate(policies, start=1):
                    row = policies_table.add_row()
                    # Ref column: same style as Process Description Ref (Avenir, 14, bold, red, centered)
                    ref_cell = row.cells[0]
                    ref_cell.text = ''
                    ref_para = ref_cell.paragraphs[0]
                    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = ref_para.add_run(str(idx))
                    run.font.name = 'Avenir LT Std 45 Book'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    # Policy text column: Avenir, 12, bold
                    text_cell = row.cells[1]
                    text_cell.text = ''
                    text_para = text_cell.paragraphs[0]
                    run = text_para.add_run(policy.get('policy', ''))
                    run.font.name = 'Avenir LT Std 45 Book'
                    run.font.size = Pt(12)
                    run.font.bold = True
            else:
                row = policies_table.add_row()
                # N/A ref cell
                ref_cell = row.cells[0]
                ref_cell.text = ''
                ref_para = ref_cell.paragraphs[0]
                ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = ref_para.add_run('N/A')
                run.font.name = 'Avenir LT Std 45 Book'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)
                # N/A policy text cell
                text_cell = row.cells[1]
                text_cell.text = ''
                text_para = text_cell.paragraphs[0]
                run = text_para.add_run('N/A')
                run.font.name = 'Avenir LT Std 45 Book'
                run.font.size = Pt(12)
                run.font.bold = True

        # Save to BytesIO
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    except Exception as e:
        print(f"[ERROR] Word Doc Generation Failed: {e}")
        import traceback
        traceback.print_exc()
        return None

@app.route('/')
def index():
    try:
        debug_log(f"index() called, template_folder={app.template_folder}")
        return render_template('index.html')
    except Exception as e:
        import traceback
        debug_log(f"index() ERROR: {e}")
        debug_log(traceback.format_exc())
        raise

@app.route('/extract-metadata', methods=['POST'])
def extract_metadata():
    """Extract metadata from uploaded BPMN file or pasted XML for form auto-population"""
    bpmn_content = None

    # Check for file upload first
    if 'bpmn_file' in request.files:
        file = request.files['bpmn_file']
        if file.filename != '':
            bpmn_content = file.read()

    # Fall back to raw XML text
    if not bpmn_content:
        xml_code = request.form.get('xml_code', '').strip()
        if xml_code:
            bpmn_content = xml_code.encode('utf-8')

    if not bpmn_content:
        return jsonify({'error': 'No BPMN content provided'}), 400

    try:
        metadata = extract_metadata_from_bpmn(bpmn_content)
        return jsonify({'success': True, 'metadata': metadata})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/history', methods=['GET'])
def get_history():
    """Get all history entries"""
    history_manager.set_user('local')
    history = history_manager.get_all()
    return jsonify(history)

@app.route('/api/history/<int:index>', methods=['GET'])
def get_history_entry(index):
    """Get a specific history entry by index"""
    history_manager.set_user('local')
    entry = history_manager.get_entry(index)
    if entry:
        return jsonify(entry)
    return jsonify({'error': 'Entry not found'}), 404

# --- Archive API Endpoints ---

@app.route('/api/user/set', methods=['POST'])
def set_user():
    """Set current user ID in session"""
    data = request.get_json()
    user_id = data.get('user_id', '').strip().lower()  # Normalize to lowercase

    if not user_id:
        return jsonify({'error': 'User ID required'}), 400

    # Simple validation - alphanumeric and underscores only
    if not re.match(r'^[a-zA-Z0-9_]+$', user_id):
        return jsonify({'error': 'User ID can only contain letters, numbers, and underscores'}), 400

    session['user_id'] = user_id
    return jsonify({'success': True, 'user_id': user_id})

@app.route('/api/user/get', methods=['GET'])
def get_user():
    """Get current user ID from session"""
    user_id = session.get('user_id', None)
    return jsonify({'user_id': user_id})

@app.route('/api/archive/save', methods=['POST'])
def save_archive():
    """Save current BPMN and Word files to archive"""
    user_id = session.get('user_id', None)
    if not user_id:
        return jsonify({'error': 'No user logged in'}), 401

    # Get files from request
    if 'bpmn_file' not in request.files or 'docx_file' not in request.files:
        return jsonify({'error': 'Both BPMN and Word files required'}), 400

    bpmn_file = request.files['bpmn_file']
    docx_file = request.files['docx_file']
    process_name = request.form.get('process_name', 'Untitled Process')

    if bpmn_file.filename == '' or docx_file.filename == '':
        return jsonify({'error': 'Empty file names'}), 400

    # Save files temporarily
    import tempfile
    with tempfile.NamedTemporaryFile(delete=False, suffix='.bpmn') as tmp_bpmn:
        bpmn_file.save(tmp_bpmn.name)
        bpmn_temp_path = tmp_bpmn.name

    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
        docx_file.save(tmp_docx.name)
        docx_temp_path = tmp_docx.name

    try:
        # Save to archive
        archive_id = archive_manager.save_archive(
            user_id=user_id,
            process_name=process_name,
            bpmn_file_path=bpmn_temp_path,
            docx_file_path=docx_temp_path
        )

        return jsonify({
            'success': True,
            'archive_id': archive_id,
            'message': 'Files archived successfully'
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up temp files
        try:
            os.remove(bpmn_temp_path)
            os.remove(docx_temp_path)
        except:
            pass

@app.route('/api/archive/list', methods=['GET'])
def list_archives():
    """Get all archives"""
    archives = archive_manager.get_user_archives('local')
    return jsonify({'archives': archives})

@app.route('/api/archive/<int:archive_id>/bpmn', methods=['GET'])
def download_archive_bpmn(archive_id):
    """Download archived BPMN file"""
    archive = archive_manager.get_archive(archive_id)
    if not archive:
        return jsonify({'error': 'Archive not found'}), 404

    bpmn_path = archive_manager.get_file_path(archive_id, 'bpmn')
    if not bpmn_path:
        return jsonify({'error': 'BPMN file not found'}), 404

    # Use save dialog for download
    import tkinter as tk
    from tkinter import filedialog
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    save_path = filedialog.asksaveasfilename(
        parent=root,
        defaultextension='.bpmn',
        filetypes=[('BPMN File', '*.bpmn')],
        initialfile=archive['bpmn_filename'],
        title='Save BPMN File'
    )
    root.destroy()

    if save_path:
        import shutil
        shutil.copy(bpmn_path, save_path)
        return jsonify({'success': True, 'path': save_path})
    return jsonify({'success': False, 'message': 'Cancelled'})

@app.route('/api/archive/<int:archive_id>/docx', methods=['GET'])
def download_archive_docx(archive_id):
    """Download archived Word file"""
    archive = archive_manager.get_archive(archive_id)
    if not archive:
        return jsonify({'error': 'Archive not found'}), 404

    docx_path = archive_manager.get_file_path(archive_id, 'docx')
    if not docx_path:
        return jsonify({'error': 'Word file not found'}), 404

    # Use save dialog for download
    import tkinter as tk
    from tkinter import filedialog
    root = tk.Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    save_path = filedialog.asksaveasfilename(
        parent=root,
        defaultextension='.docx',
        filetypes=[('Word Document', '*.docx')],
        initialfile=archive['docx_filename'],
        title='Save Word Document'
    )
    root.destroy()

    if save_path:
        import shutil
        shutil.copy(docx_path, save_path)
        return jsonify({'success': True, 'path': save_path})
    return jsonify({'success': False, 'message': 'Cancelled'})

@app.route('/api/archive/<int:archive_id>', methods=['DELETE'])
def delete_archive(archive_id):
    """Delete an archive"""
    success = archive_manager.delete_archive(archive_id, 'local')
    if success:
        return jsonify({'success': True, 'message': 'Archive deleted'})
    return jsonify({'error': 'Archive not found'}), 404

@app.route('/api/generate-from-xml', methods=['POST'])
def api_generate_from_xml():
    """API endpoint for Camunda Modeler plugin - accepts XML, returns .docx"""
    try:
        data = request.get_json()
        if not data or 'xml' not in data:
            return jsonify({'error': 'No XML provided'}), 400

        bpmn_content = data['xml'].encode('utf-8')
        metadata = data.get('metadata', {})

        # Extract BPMN metadata for fields not provided
        bpmn_metadata = extract_metadata_from_bpmn(bpmn_content)
        for field in ['process_name', 'process_code', 'purpose', 'scope']:
            if not metadata.get(field, '').strip() and field in bpmn_metadata:
                metadata[field] = bpmn_metadata[field]

        # Auto-populate abbreviations if not provided
        if 'abbreviations_list' not in metadata:
            metadata['abbreviations_list'] = bpmn_metadata.get('abbreviations_list', [])

        # Auto-populate references if not provided
        if 'references_list' not in metadata:
            references = []
            lane_names = bpmn_metadata.get('lane_names', [])
            for lane_name in lane_names:
                references.append({'id': 'N/A', 'title': f"{lane_name} Approval"})
            process_code = metadata.get('process_code', bpmn_metadata.get('process_code', ''))
            process_name = metadata.get('process_name', bpmn_metadata.get('process_name', ''))
            if process_code or process_name:
                diagram_id = f"DGM- {process_code}" if process_code else "DGM-"
                diagram_title = f"{process_name} Process Diagram        Notations Meaning"
                references.append({'id': diagram_id, 'title': diagram_title})
            metadata['references_list'] = references

        # Auto-populate general policies if not provided
        if 'general_policies_list' not in metadata:
            metadata['general_policies_list'] = bpmn_metadata.get('general_policies_list', [])

        template_name = (data.get('template') or data.get('sop_template') or '').strip() or None
        context = parse_bpmn_to_context(bpmn_content, metadata)
        file_stream = create_word_doc_from_template(context, template_name=template_name)

        if not file_stream:
            return jsonify({'error': 'Failed to generate document'}), 500

        file_stream.seek(0)
        output_name = metadata.get('process_name', bpmn_metadata.get('process_name', 'SOP_Document'))
        response = make_response(file_stream.read())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{output_name}.docx"'
        response.headers['Access-Control-Allow-Origin'] = '*'
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

import uuid as _uuid
_preview_sessions = {}

@app.route('/api/upload-xml', methods=['POST', 'OPTIONS'])
def api_upload_xml():
    """Store XML temporarily and return a session ID for the preview page"""
    if request.method == 'OPTIONS':
        response = make_response()
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return response

    data = request.get_json()
    if not data or 'xml' not in data:
        return jsonify({'error': 'No XML provided'}), 400

    session_id = str(_uuid.uuid4())
    bpmn_content = data['xml'].encode('utf-8')
    metadata = extract_metadata_from_bpmn(bpmn_content)

    _preview_sessions[session_id] = {
        'xml': data['xml'],
        'metadata': metadata
    }

    resp = jsonify({'session_id': session_id, 'metadata': metadata})
    resp.headers['Access-Control-Allow-Origin'] = '*'
    return resp

@app.route('/preview/<session_id>')
def preview_page(session_id):
    """Serve the SOP preview/edit form pre-populated from BPMN"""
    session_data = _preview_sessions.get(session_id)
    if not session_data:
        return 'Session expired. Please try again from Camunda Modeler.', 404

    meta = session_data['metadata']
    return render_template('preview.html', session_id=session_id, meta=meta, available_templates=ALLOWED_SOP_TEMPLATES)

@app.route('/api/generate-and-download/<session_id>', methods=['POST', 'OPTIONS'])
def api_generate_and_download(session_id):
    """Generate .docx from stored XML + user-edited metadata, return binary"""
    if request.method == 'OPTIONS':
        response = make_response()
        response.headers['Access-Control-Allow-Origin'] = '*'
        response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
        return response

    session_data = _preview_sessions.get(session_id)
    if not session_data:
        resp = jsonify({'error': 'Session expired'})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 404

    try:
        bpmn_content = session_data['xml'].encode('utf-8')
        form_data = request.form.to_dict()

        metadata = {}
        for field in ['process_name', 'process_code', 'issued_by', 'release_date', 'process_owner', 'purpose', 'scope']:
            metadata[field] = form_data.get(field, '').strip()

        # Fall back to BPMN metadata for empty fields
        bpmn_meta = session_data['metadata']
        for field in ['process_name', 'process_code', 'purpose', 'scope']:
            if not metadata.get(field) and field in bpmn_meta:
                metadata[field] = bpmn_meta[field]

        # Parse abbreviations
        abbrev_terms = request.form.getlist('abbrev_term[]')
        abbrev_defs = request.form.getlist('abbrev_def[]')
        abbreviations = []
        for term, definition in zip(abbrev_terms, abbrev_defs):
            if term.strip() or definition.strip():
                abbreviations.append({'term': term.strip(), 'definition': definition.strip()})
        if not abbreviations and 'abbreviations_list' in bpmn_meta:
            abbreviations = bpmn_meta['abbreviations_list']
        metadata['abbreviations_list'] = abbreviations

        # Parse references
        ref_ids = request.form.getlist('ref_id[]')
        ref_titles = request.form.getlist('ref_title[]')
        references = []
        for doc_id, title in zip(ref_ids, ref_titles):
            if doc_id.strip() or title.strip():
                references.append({'id': doc_id.strip(), 'title': title.strip()})
        if not references:
            lane_names = bpmn_meta.get('lane_names', [])
            for lane_name in lane_names:
                references.append({'id': 'N/A', 'title': f"{lane_name} Approval"})
            process_code = metadata.get('process_code', '')
            process_name = metadata.get('process_name', '')
            if process_code or process_name:
                references.append({
                    'id': f"DGM- {process_code}" if process_code else "DGM-",
                    'title': f"{process_name} Process Diagram        Notations Meaning"
                })
        metadata['references_list'] = references

        # Parse policies
        policy_refs = request.form.getlist('policy_ref[]')
        policy_texts = request.form.getlist('policy_text[]')
        policies = []
        for ref, text in zip(policy_refs, policy_texts):
            if ref.strip() or text.strip():
                policies.append({'ref': ref.strip(), 'policy': text.strip()})
        metadata['general_policies_list'] = policies

        template_name = (request.form.get('sop_template') or '').strip() or None
        context = parse_bpmn_to_context(bpmn_content, metadata)
        file_stream = create_word_doc_from_template(context, template_name=template_name)

        if not file_stream:
            return jsonify({'error': 'Failed to generate document'}), 500

        file_stream.seek(0)
        output_name = metadata.get('process_name', 'SOP_Document')

        # Clean up session
        _preview_sessions.pop(session_id, None)

        response = make_response(file_stream.read())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{output_name}.docx"'
        response.headers['Access-Control-Allow-Origin'] = '*'
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        resp = jsonify({'error': str(e)})
        resp.headers['Access-Control-Allow-Origin'] = '*'
        return resp, 500

@app.route('/api/generate-from-xml', methods=['OPTIONS'])
def api_generate_from_xml_options():
    """CORS preflight for Camunda Modeler plugin"""
    response = make_response()
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type'
    return response

@app.route('/generate', methods=['POST'])
def generate_sop():
    input_type = request.form.get('input_type')
    bpmn_content = None
    output_name = None

    if input_type == 'bpmn':
        if 'bpmn_file' not in request.files:
            return "No file part", 400
        file = request.files['bpmn_file']
        if file.filename == '':
            return "No selected file", 400
        if file:
            bpmn_content = file.read()
            # Use BPMN filename (without extension) for output
            output_name = file.filename.rsplit('.', 1)[0] if '.' in file.filename else file.filename
    elif input_type == 'xml':
        xml_code = request.form.get('xml_code')
        if not xml_code:
            return "No XML code provided", 400
        bpmn_content = xml_code.encode('utf-8') # Encode to bytes for consistency with file.read()

        # Extract pool/process name from XML
        try:
            from lxml import etree
            root = etree.fromstring(bpmn_content)
            ns = {'bpmn': 'http://www.omg.org/spec/BPMN/20100524/MODEL'}

            # Try to get participant (pool) name
            participant = root.find('.//bpmn:participant', namespaces=ns)
            if participant is not None and participant.get('name'):
                output_name = participant.get('name')
            else:
                # Fallback to process name
                process = root.find('.//bpmn:process', namespaces=ns)
                if process is not None and process.get('name'):
                    output_name = process.get('name')
        except:
            pass
    else:
        return "Invalid input type selected", 400

    if bpmn_content:
        metadata = request.form.to_dict()

        # Extract BPMN metadata for fields user left empty
        bpmn_metadata = extract_metadata_from_bpmn(bpmn_content)

        # For simple string fields: use form value if non-empty, else BPMN value
        for field in ['process_name', 'process_code', 'purpose', 'scope']:
            form_value = metadata.get(field, '').strip()
            if not form_value and field in bpmn_metadata:
                metadata[field] = bpmn_metadata[field]

        # Parse abbreviation entries
        abbrev_terms = request.form.getlist('abbrev_term[]')
        abbrev_defs = request.form.getlist('abbrev_def[]')
        abbreviations = []
        for term, definition in zip(abbrev_terms, abbrev_defs):
            if term.strip() or definition.strip():  # Only add non-empty entries
                abbreviations.append({'term': term.strip(), 'definition': definition.strip()})
        # If no user abbreviations, fall back to BPMN-extracted ones
        if not abbreviations and 'abbreviations_list' in bpmn_metadata:
            abbreviations = bpmn_metadata['abbreviations_list']
        metadata['abbreviations_list'] = abbreviations

        # Parse reference document entries
        ref_ids = request.form.getlist('ref_id[]')
        ref_titles = request.form.getlist('ref_title[]')
        references = []
        for doc_id, title in zip(ref_ids, ref_titles):
            if doc_id.strip() or title.strip():  # Only add non-empty entries
                references.append({'id': doc_id.strip(), 'title': title.strip()})

        # Auto-add lane approvals + DGM row ONLY if form sent no references
        # (i.e., JavaScript auto-fill didn't run). If user already has references
        # from auto-fill (possibly edited), don't add duplicates.
        if not references:
            lane_names = bpmn_metadata.get('lane_names', [])
            for lane_name in lane_names:
                references.append({'id': 'N/A', 'title': f"{lane_name} Approval"})

            process_code = metadata.get('process_code', '').strip()
            if not process_code:
                process_code = bpmn_metadata.get('process_code', '')
            process_name = metadata.get('process_name', '').strip()
            if not process_name:
                process_name = bpmn_metadata.get('process_name', '')
            if process_code or process_name:
                diagram_id = f"DGM- {process_code}" if process_code else "DGM-"
                diagram_title = f"{process_name} Process Diagram        Notations Meaning"
                references.append({'id': diagram_id, 'title': diagram_title})

        metadata['references_list'] = references

        # Parse general policy entries
        policy_refs = request.form.getlist('policy_ref[]')
        policy_texts = request.form.getlist('policy_text[]')
        policies = []
        for ref, text in zip(policy_refs, policy_texts):
            if ref.strip() or text.strip():
                policies.append({'ref': ref.strip(), 'policy': text.strip()})
        metadata['general_policies_list'] = policies

        template_name = (request.form.get('sop_template') or '').strip() or None
        context = parse_bpmn_to_context(bpmn_content, metadata)
        file_stream = create_word_doc_from_template(context, template_name=template_name)
        if file_stream:
            # Save to history on successful generation
            history_manager.set_user('local')
            history_data = {
                'process_name': metadata.get('process_name', ''),
                'process_code': metadata.get('process_code', ''),
                'purpose': metadata.get('purpose', ''),
                'scope': metadata.get('scope', ''),
                'abbreviations_list': metadata.get('abbreviations_list', []),
                'references_list': metadata.get('references_list', []),
                'general_policies_list': metadata.get('general_policies_list', [])
            }
            history_manager.add_entry(history_data)

            # Use extracted name or fallback
            if not output_name:
                output_name = metadata.get('process_name', 'Generated')

            # Show save dialog and save file
            import tkinter as tk
            from tkinter import filedialog

            # Create hidden root window for dialog
            root = tk.Tk()
            root.withdraw()
            root.attributes('-topmost', True)

            # Show save dialog
            save_path = filedialog.asksaveasfilename(
                parent=root,
                defaultextension='.docx',
                filetypes=[('Word Document', '*.docx')],
                initialfile=f"{output_name}.docx",
                title='Save SOP Document'
            )

            root.destroy()

            if save_path:
                # Save the file
                file_stream.seek(0)
                with open(save_path, 'wb') as f:
                    f.write(file_stream.read())

                # Auto-save to archive
                import tempfile
                try:
                    # Save BPMN to temp file
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.bpmn') as tmp_bpmn:
                        tmp_bpmn.write(bpmn_content)
                        tmp_bpmn_path = tmp_bpmn.name

                    # Save DOCX to temp file
                    file_stream.seek(0)
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                        tmp_docx.write(file_stream.read())
                        tmp_docx_path = tmp_docx.name

                    # Archive the files
                    archive_manager.save_archive(
                        user_id='local',
                        process_name=output_name,
                        bpmn_file_path=tmp_bpmn_path,
                        docx_file_path=tmp_docx_path
                    )

                    # Clean up temp files
                    os.remove(tmp_bpmn_path)
                    os.remove(tmp_docx_path)
                except Exception as e:
                    debug_log(f"Error archiving: {e}")

                return jsonify({'success': True, 'message': f'Document saved to {save_path}', 'path': save_path})
            else:
                return jsonify({'success': False, 'message': 'Save cancelled'})
    return "An error occurred during file processing. Check the console for details.", 500

def start_server():
    """Start Flask server in background thread"""
    serve(app, host='127.0.0.1', port=8000, _quiet=True)

# Global reference to webview window for save dialogs
webview_window = None

def main():
    """Main entry point - launches native window with pywebview"""
    global webview_window
    import webview

    # Start Flask server in background thread
    server_thread = threading.Thread(target=start_server, daemon=True)
    server_thread.start()

    # Give server time to start
    import time
    time.sleep(0.5)

    # Create native window
    webview.create_window(
        'BPMN to SOP Generator',
        'http://127.0.0.1:8000',
        width=1000,
        height=800,
        resizable=True,
        min_size=(800, 600)
    )
    webview.start()

if __name__ == '__main__':
    main()